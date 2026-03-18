"""
CREAR_HIMNARIO.PY
=================
Genera un documento Word (DOCX) profesional con formato de 2 columnas
desde el archivo Himnario_Limpio_Final.txt

Características:
- Layout de 2 columnas con espaciado de 720 twips
- Previene cortes de estrofas entre columnas (keep-with-next)
- Títulos en negrita 10pt
- Contenido en 9pt
- Coros en negrita
- Detección de CORO I, II, III, IV

Uso:
    python crear_himnario.py
"""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_keep_with_next(paragraph):
    """Previene que el párrafo se separe del siguiente"""
    pPr = paragraph._element.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    pPr.append(keepNext)

def set_keep_lines_together(paragraph):
    """Mantiene todas las líneas del párrafo juntas"""
    pPr = paragraph._element.get_or_add_pPr()
    keepLines = OxmlElement('w:keepLines')
    pPr.append(keepLines)

def es_marcador_coro(linea):
    """Detecta si una línea es CORO, ESTRIBILLO o FINAL con numeración romana"""
    patron = r'^\s*(coro|estribillo|final)(\s+[ivx]+)?\s*:?\s*$'
    return bool(re.match(patron, linea.strip(), re.IGNORECASE))

def agrupar_en_estrofas(lineas):
    """Agrupa líneas en estrofas separadas por líneas vacías"""
    estrofas = []
    estrofa_actual = []
    
    for linea in lineas:
        if linea.strip():
            estrofa_actual.append(linea)
        else:
            if estrofa_actual:
                estrofas.append(estrofa_actual)
                estrofa_actual = []
    
    if estrofa_actual:
        estrofas.append(estrofa_actual)
    
    return estrofas

def crear_documento():
    """Crea el documento Word con configuración de 2 columnas"""
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        # Configurar 2 columnas con espacio de 720 twips
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '720')
        if not sectPr.xpath('./w:cols'):
            sectPr.append(cols)
    
    return doc

def agregar_portada(doc):
    """Agrega la portada del himnario - diseño V1"""
    # Espacios superiores
    for _ in range(5):
        doc.add_paragraph()
    
    # Título principal
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = 1  # Centro
    run = p_titulo.add_run("HIMNARIO ESPERANZA V1")
    run.bold = True
    run.font.size = Pt(48)  # 609600 twips = 48pt
    
    # Espacios
    for _ in range(4):
        doc.add_paragraph()
    
    # Subtítulo 1
    p_sub1 = doc.add_paragraph()
    p_sub1.alignment = 1
    run = p_sub1.add_run("Colección de Himnos Cristianos y del Mensaje")
    run.font.size = Pt(18)  # 228600 twips ≈ 18pt
    
    # Subtítulo 2
    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = 1
    run = p_sub2.add_run("Tabernáculo La Esperanza")
    run.font.size = Pt(18)
    
    # Subtítulo 3
    p_sub3 = doc.add_paragraph()
    p_sub3.alignment = 1
    run = p_sub3.add_run("Calarcá - Quindío")
    run.font.size = Pt(18)
    
    doc.add_paragraph()
    
    # Año
    p_ano = doc.add_paragraph()
    p_ano.alignment = 1
    run = p_ano.add_run("2025")
    run.font.size = Pt(14)  # 177800 twips ≈ 14pt
    
    # Salto de página
    doc.add_page_break()

def agregar_indice(doc, himnos):
    """Agrega el índice alfabético - diseño V1"""
    # Espacios superiores
    for _ in range(3):
        doc.add_paragraph()
    
    # Título del índice
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = 1
    run = p_titulo.add_run("ÍNDICE")
    run.bold = True
    run.font.size = Pt(16)  # 203200 twips ≈ 16pt
    
    doc.add_paragraph()
    
    # Agregar himnos al índice (sin ordenar, en orden numérico)
    for numero, titulo, _ in himnos:
        p = doc.add_paragraph()
        run = p.add_run(f"{numero}. {titulo}")
        run.font.size = Pt(9)  # 114300 twips ≈ 9pt
    
    # Salto de página
    doc.add_page_break()

def agregar_himno(doc, numero, titulo, letra):
    """Agrega un himno al documento con formato apropiado"""
    # Agregar título
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run(f"{numero}. {titulo}")
    run_titulo.bold = True
    run_titulo.font.size = Pt(10)
    set_keep_with_next(p_titulo)
    
    # Procesar letra por estrofas
    lineas = letra.strip().split('\n')
    estrofas = agrupar_en_estrofas(lineas)
    
    for i, estrofa in enumerate(estrofas):
        # Procesar cada línea de la estrofa
        for j, linea in enumerate(estrofa):
            p = doc.add_paragraph()
            
            # Verificar si es CORO
            if es_marcador_coro(linea):
                run = p.add_run(linea.strip())
                run.bold = True
                run.font.size = Pt(9)
            else:
                run = p.add_run(linea.strip())
                run.font.size = Pt(9)
            
            # Aplicar keep-with-next a todas las líneas excepto la última de la estrofa
            if j < len(estrofa) - 1:
                set_keep_with_next(p)
                set_keep_lines_together(p)
        
        # Espacio entre estrofas (no después de la última)
        if i < len(estrofas) - 1:
            doc.add_paragraph()
    
    # Doble línea vacía entre himnos
    doc.add_paragraph()
    doc.add_paragraph()

def main():
    """Función principal"""
    print('=' * 80)
    print('GENERADOR DE HIMNARIO WORD (DOCX)')
    print('=' * 80)
    
    # Leer archivo fuente
    print('\n[1/5] Leyendo Himnario_Limpio_Final.txt de post-proceso/...')
    import os
    post_dir = os.path.join(os.path.dirname(__file__), '..', 'post-proceso')
    archivo_entrada = os.path.join(post_dir, 'Himnario_Limpio_Final.txt')
    archivo_salida = os.path.join(post_dir, 'Himnario_TabernaculoLaEsperanza_V1.docx')
    with open(archivo_entrada, 'r', encoding='utf-8') as f:
        contenido = f.read()
    
    # Parsear himnos
    print('[2/5] Procesando himnos...')
    patron = r'(\d+)\.\s+(.+?)\n((?:(?!\n\d+\.\s).)+)'
    himnos = re.findall(patron, contenido, re.DOTALL)
    print(f'      Total de himnos encontrados: {len(himnos)}')
    
    # Crear documento
    print('[3/5] Creando portada...')
    doc = crear_documento()
    agregar_portada(doc)
    
    print('[4/5] Generando índice alfabético...')
    agregar_indice(doc, himnos)
    
    # Agregar himnos
    print('[5/5] Agregando himnos al documento...')
    for numero, titulo, letra in himnos:
        agregar_himno(doc, numero, titulo.strip(), letra)
    
    # Guardar documento
    doc.save(archivo_salida)
    
    print('\n' + '=' * 80)
    print(f'OK Documento generado exitosamente: {archivo_salida}')
    print('OK Portada incluida')
    print('OK Índice incluido (orden numérico)')
    print(f'OK Total de himnos: {len(himnos)}')
    print('OK Formato: 2 columnas con protección anti-corte de estrofas')
    print('=' * 80)

if __name__ == '__main__':
    main()
